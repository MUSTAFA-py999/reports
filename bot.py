import os
import asyncio
import threading
import logging
import html as html_lib
from flask import Flask
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import (
    ApplicationBuilder, ContextTypes, CommandHandler,
    MessageHandler, CallbackQueryHandler, filters
)
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.output_parsers import PydanticOutputParser
from langchain_core.messages import HumanMessage
from pydantic import BaseModel, Field
from typing import List, Optional
from io import BytesIO
from weasyprint import HTML as WeasyHTML
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO
)
logger = logging.getLogger(__name__)

flask_app = Flask(__name__)

@flask_app.route('/')
def home():
    return "✅ Repooreto Bot v5.2"

@flask_app.route('/health')
def health():
    return {"status": "healthy", "version": "5.2"}, 200

def run_flask():
    port = int(os.environ.get("PORT", 10000))
    flask_app.run(host='0.0.0.0', port=port, debug=False, use_reloader=False)


# ═══════════════════════════════════════════════════
# QUEUE SYSTEM
# ═══════════════════════════════════════════════════
report_queue: asyncio.Queue = None
active_jobs    = {}
queue_positions = {}
MAX_CONCURRENT  = 2


async def queue_worker(app):
    semaphore = asyncio.Semaphore(MAX_CONCURRENT)

    async def process_one(user_id, session, msg_id):
        async with semaphore:
            active_jobs[user_id] = True
            for uid in list(queue_positions.keys()):
                if queue_positions[uid] > 0:
                    queue_positions[uid] -= 1
            try:
                loop = asyncio.get_event_loop()
                # اختيار الصيغة المطلوبة
                file_format = session.get("file_format", "pdf")
                if file_format == "pdf":
                    pdf_bytes, title = await loop.run_in_executor(None, generate_report, session)
                    file_bytes = pdf_bytes
                    extension = "pdf"
                else:
                    docx_bytes, title = await loop.run_in_executor(None, generate_word_report, session)
                    file_bytes = docx_bytes
                    extension = "docx"

                lang_name  = LANGUAGES[session.get("language", "ar")]["name"]
                depth_name = DEPTH_OPTIONS[session.get("depth", "medium")]["name"]
                tpl_name   = "🎨 مخصص" if session.get("custom_mode") else TEMPLATES.get(session.get("template", "emerald"), {}).get("name", "")

                if file_bytes:
                    safe_name  = "".join(c if c.isalnum() or c in (' ', '_', '-') else '_' for c in title[:40])
                    safe_title = title.replace('<','&lt;').replace('>','&gt;').replace('&','&amp;')
                    caption = (
                        f"👻 <b>تقريرك جاهز يا طالبنا!</b>\n\n"
                        f"📄 <b>{safe_title}</b>\n"
                        f"🌐 {lang_name}  |  📏 {depth_name}  |  🎨 {tpl_name}\n"
                        f"📁 الصيغة: {file_format.upper()}\n\n"
                        f"🔄 أرسل موضوعاً جديداً لتقرير آخر!"
                    )
                    await app.bot.send_document(
                        chat_id=user_id,
                        document=BytesIO(file_bytes),
                        filename=f"{safe_name}.{extension}",
                        caption=caption,
                        parse_mode='HTML'
                    )
                    try:
                        await app.bot.delete_message(chat_id=user_id, message_id=msg_id)
                    except Exception:
                        pass
                    logger.info(f"✅ Report sent to {user_id} as {extension}")
                else:
                    err = str(title).replace('<','&lt;').replace('>','&gt;').replace('&','&amp;')
                    await app.bot.send_message(
                        chat_id=user_id,
                        text=f"❌ <b>فشل إنشاء التقرير:</b>\n{err[:300]}\n\n🔄 أرسل موضوعاً جديداً.",
                        parse_mode='HTML'
                    )
            except Exception as e:
                logger.error(f"Queue worker error for {user_id}: {e}", exc_info=True)
                err = str(e)[:200].replace('<','&lt;').replace('>','&gt;').replace('&','&amp;')
                await app.bot.send_message(
                    chat_id=user_id,
                    text=f"❌ <b>خطأ غير متوقع:</b>\n<code>{err}</code>\n\n🔄 أرسل موضوعاً جديداً.",
                    parse_mode='HTML'
                )
            finally:
                active_jobs.pop(user_id, None)
                queue_positions.pop(user_id, None)
                user_sessions.pop(user_id, None)

    while True:
        item = await report_queue.get()
        user_id, session, msg_id = item
        asyncio.create_task(process_one(user_id, session, msg_id))
        report_queue.task_done()


# ═══════════════════════════════════════════════════
# PYDANTIC MODELS
# ═══════════════════════════════════════════════════
class SmartQuestions(BaseModel):
    questions: List[str] = Field(
        description="List of open-ended questions (2-5) to ask the student about their report."
    )

class ReportBlock(BaseModel):
    block_type: str = Field(
        description=(
            "Block type — ONE of: 'paragraph','bullets','numbered_list',"
            "'table','pros_cons','comparison','stats','examples','quote'"
        )
    )
    title: str = Field(description="Section heading")
    style: Optional[str]               = Field(default=None, description="pros_cons style: A/B/C/D")
    text:  Optional[str]               = Field(default=None)
    items: Optional[List[str]]         = Field(default=None)
    pros:  Optional[List[str]]         = Field(default=None)
    cons:  Optional[List[str]]         = Field(default=None)
    headers:      Optional[List[str]]  = Field(default=None)
    rows:         Optional[List[List[str]]] = Field(default=None)
    side_a:       Optional[str]        = Field(default=None)
    side_b:       Optional[str]        = Field(default=None)
    criteria:     Optional[List[str]]  = Field(default=None)
    side_a_values:Optional[List[str]]  = Field(default=None)
    side_b_values:Optional[List[str]]  = Field(default=None)

class DynamicReport(BaseModel):
    title:        str              = Field(description="Report title")
    introduction: str              = Field(description="Introduction: 3-5 sentences. Direct and engaging.")
    blocks:       List[ReportBlock]= Field(description="Content blocks")
    conclusion:   str              = Field(description="Conclusion: 4-6 sentences. Genuine insight. MANDATORY.")


# ═══════════════════════════════════════════════════
# CONFIG
# ═══════════════════════════════════════════════════
user_sessions = {}

LANGUAGES = {
    "ar": {
        "name": "🇸🇦 العربية",
        "dir": "rtl", "align": "right", "lang_attr": "ar",
        "font": "'Traditional Arabic', 'Arial', sans-serif",
        "intro_label": "المقدمة",
        "conclusion_label": "الخاتمة",
        "pros_label": "✅ المزايا",
        "cons_label": "❌ العيوب",
        "instruction": "Write ALL content in formal Arabic (فصحى). Every word must be Arabic.",
        "q_prompt": (
            "أنت مساعد أكاديمي لطلاب الجامعة.\n"
            "الطالب يريد تقريراً عن: \"{topic}\".\n\n"
            "اكتب بالعربية 2-4 أسئلة قصيرة ومباشرة لتحديد ما يريده الطالب في تقريره.\n"
            "قواعد الأسئلة:\n"
            "- قصيرة (جملة واحدة فقط لكل سؤال)\n"
            "- مباشرة ومحددة\n"
            "- موضوعات بسيطة: 2 أسئلة — معقدة: 3-4 أسئلة\n"
        ),
    },
    "en": {
        "name": "🇬🇧 English",
        "dir": "ltr", "align": "left", "lang_attr": "en",
        "font": "'Arial', 'Helvetica', sans-serif",
        "intro_label": "Introduction",
        "conclusion_label": "Conclusion",
        "pros_label": "✅ Pros",
        "cons_label": "❌ Cons",
        "instruction": "Write ALL content in English. Every word must be English.",
        "q_prompt": (
            "أنت مساعد أكاديمي لطلاب الجامعة.\n"
            "الطالب يريد تقريراً إنجليزياً عن: \"{topic}\".\n\n"
            "اكتب بالعربية 2-4 أسئلة قصيرة ومباشرة لتحديد ما يريده الطالب في تقريره.\n"
            "قواعد الأسئلة:\n"
            "- قصيرة (جملة واحدة فقط لكل سؤال)\n"
            "- مباشرة ومحددة\n"
            "- موضوعات بسيطة: 2 أسئلة — معقدة: 3-4 أسئلة\n"
        ),
    },
}

# ── قوالب جاهزة ──
TEMPLATES = {
    "emerald":      {"name": "🌿 زمردي",     "primary": "#1a4731", "accent": "#52b788", "bg": "#f0faf4", "bg2": "#ffffff"},
    "modern":       {"name": "🚀 عصري",      "primary": "#5a67d8", "accent": "#667eea", "bg": "#ebf4ff", "bg2": "#ffffff"},
    "minimal":      {"name": "⚪ بسيط",      "primary": "#2d3748", "accent": "#718096", "bg": "#f7fafc", "bg2": "#ffffff"},
    "professional": {"name": "💼 احترافي",   "primary": "#1a365d", "accent": "#2b6cb0", "bg": "#bee3f8", "bg2": "#f0f4ff"},
    "dark_elegant": {"name": "🖤 أنيق داكن", "primary": "#d4af37", "accent": "#f6d860", "bg": "#2d3748", "bg2": "#4a5568"},
    "royal":        {"name": "👑 ملكي ذهبي", "primary": "#5b0e2d", "accent": "#c9a227", "bg": "#fdf6e3", "bg2": "#fff9f0"},
}

# ── خيارات التخصيص ──
CUSTOM_COLORS = {
    "royal_blue": {"label": "🔵 أزرق ملكي",    "primary": "#1a365d", "accent": "#3182ce", "bg": "#ebf8ff", "bg2": "#ffffff"},
    "emerald_g":  {"label": "🌿 زمردي",         "primary": "#1a4731", "accent": "#38a169", "bg": "#f0fff4", "bg2": "#ffffff"},
    "purple":     {"label": "💜 بنفسجي",        "primary": "#44337a", "accent": "#805ad5", "bg": "#faf5ff", "bg2": "#ffffff"},
    "orange":     {"label": "🟠 برتقالي دافئ", "primary": "#7b341e", "accent": "#dd6b20", "bg": "#fffaf0", "bg2": "#ffffff"},
    "slate":      {"label": "⚫ رمادي راقٍ",   "primary": "#1a202c", "accent": "#718096", "bg": "#f7fafc", "bg2": "#ffffff"},
    "crimson":    {"label": "🔴 أحمر كلاسيك",  "primary": "#742a2a", "accent": "#c53030", "bg": "#fff5f5", "bg2": "#ffffff"},
    "teal":       {"label": "🩵 تيل أنيق",     "primary": "#1d4044", "accent": "#2c7a7b", "bg": "#e6fffa", "bg2": "#ffffff"},
    "gold":       {"label": "✨ ذهبي ملكي",    "primary": "#5b0e2d", "accent": "#c9a227", "bg": "#fdf6e3", "bg2": "#fff9f0"},
}

# ── خيارات حجم الخط ──
CUSTOM_FONT_SIZES = {
    "xsmall": {"label": "🔹 صغير جداً (12px)", "size": "12px"},
    "small":  {"label": "🔸 صغير (14px)",      "size": "14px"},
    "medium": {"label": "🔹 متوسط (16px)",     "size": "16px"},
    "large":  {"label": "🔸 كبير (18px)",      "size": "18px"},
    "xlarge": {"label": "🔹 كبير جداً (20px)", "size": "20px"},
}

# ── خيارات نوع الخط ──
CUSTOM_FONTS = {
    "traditional": {"label": "📜 تقليدي",        "value": "'Traditional Arabic', serif"},
    "amiri":       {"label": "🕌 أميري",          "value": "'Amiri', serif"},
    "cairo":       {"label": "🏙 كايرو",          "value": "'Cairo', sans-serif"},
    "tajawal":     {"label": "✍️ تجوّل",         "value": "'Tajawal', sans-serif"},
    "arial":       {"label": "🔤 Arial",          "value": "Arial, sans-serif"},
    "georgia":     {"label": "📰 Georgia",        "value": "Georgia, serif"},
    "times":       {"label": "📋 Times New Roman","value": "'Times New Roman', serif"},
    "verdana":     {"label": "👁 Verdana",        "value": "Verdana, sans-serif"},
    "roboto":      {"label": "🤖 Roboto",         "value": "'Roboto', sans-serif"},
    "open-sans":   {"label": "✉️ Open Sans",      "value": "'Open Sans', sans-serif"},
}

# ── خيارات تباعد الأسطر ──
LINE_HEIGHTS = {
    "compact": {"label": "📏 مضغوط (1.5)", "value": "1.5"},
    "normal":  {"label": "📐 عادي (1.8)",   "value": "1.8"},
    "relaxed": {"label": "📏 واسع (2.2)",   "value": "2.2"},
}

# ── خيارات هوامش الصفحة ──
PAGE_MARGINS = {
    "small":  {"label": "🔹 ضيقة (1.5 سم)", "value": "1.5cm"},
    "medium": {"label": "🔸 متوسطة (2.5 سم)", "value": "2.5cm"},
    "large":  {"label": "🔻 واسعة (3.5 سم)", "value": "3.5cm"},
}

# ── خيارات نمط العنوان الرئيسي ──
HEADER_STYLES = {
    "classic": {"label": "🏛 كلاسيكي (أسود)", "color": "#000000", "size": "22px"},
    "colored": {"label": "🎨 ملون (حسب التصميم)", "color": "auto", "size": "24px"},
    "minimal": {"label": "⬜ بسيط (رمادي)", "color": "#4a5568", "size": "20px"},
}

# ── خيارات إظهار الترويسة والتذييل ──
SHOW_HEADER_FOOTER = {
    "yes": {"label": "✅ نعم، أظهرها", "show": True},
    "no":  {"label": "❌ لا، أخفها", "show": False},
}

# ── خيارات صيغة الملف ──
FILE_FORMATS = {
    "pdf": {"label": "📄 PDF", "emoji": "📄"},
    "word": {"label": "📝 Word (DOCX)", "emoji": "📝"},
}

DEPTH_OPTIONS = {
    "medium":   {"name": "📄 متوسط (3-4 صفحات)", "pages": 4,  "blocks_min": 5,  "blocks_max": 7},
    "detailed": {"name": "📚 مفصل (5-6 صفحات)",  "pages": 6,  "blocks_min": 7,  "blocks_max": 10},
    "extended": {"name": "📖 موسّع (7+ صفحات)",   "pages": 8,  "blocks_min": 10, "blocks_max": 14},
}

STATE_GUIDANCE = {
    "choosing_lang":        "🌐 من فضلك <b>اختر اللغة</b> من الأزرار أعلاه.",
    "generating_questions": "👻 الشبح يحلل موضوعك... انتظر لحظة.",
    "choosing_title":       "📌 من فضلك <b>اكتب عنوان التقرير</b> أو اضغط الزر لتركه للشبح.",
    "choosing_depth":       "📏 من فضلك <b>اختر عمق التقرير</b> من الأزرار أعلاه.",
    "choosing_style_mode":  "🎨 من فضلك <b>اختر طريقة التصميم</b> من الأزرار أعلاه.",
    "choosing_template":    "🎭 من فضلك <b>اختر قالباً</b> من الأزرار أعلاه.",
    "choosing_font_size":   "🔡 من فضلك <b>اختر حجم الخط</b> من الأزرار أعلاه.",
    "choosing_font":        "✍️ من فضلك <b>اختر نوع الخط</b> من الأزرار أعلاه.",
    "choosing_colors":      "🎨 من فضلك <b>اختر نظام الألوان</b> من الأزرار أعلاه.",
    "choosing_line_height": "📏 من فضلك <b>اختر تباعد الأسطر</b> من الأزرار أعلاه.",
    "choosing_page_margin": "📐 من فضلك <b>اختر هوامش الصفحة</b> من الأزرار أعلاه.",
    "choosing_header_style":"🎯 من فضلك <b>اختر نمط العنوان الرئيسي</b> من الأزرار أعلاه.",
    "choosing_show_header": "📰 من فضلك <b>اختر إظهار الترويسة والتذييل</b> من الأزرار أعلاه.",
    "choosing_file_format": "📁 من فضلك <b>اختر صيغة الملف</b> من الأزرار أعلاه.",
    "asking_comparison":    "📊 من فضلك <b>اختر</b> من الأزرار أعلاه.",
    "entering_comparison":  "✏️ اكتب الشيئين اللذين تريد مقارنتهما.\nمثال: <code>Python مقابل Java</code>",
    "in_queue":             "👻 تقريرك في الطابور... أرسل /cancel لإلغاء.",
}


# ═══════════════════════════════════════════════════
# LLM HELPERS
# ═══════════════════════════════════════════════════
def get_llm():
    api_key = os.getenv("GOOGLE_API_KEY")
    if not api_key:
        raise Exception("GOOGLE_API_KEY not set")
    return ChatGoogleGenerativeAI(
        model="gemini-2.5-flash",
        temperature=0.5,
        google_api_key=api_key,
        max_retries=3
    )


def generate_dynamic_questions(topic: str, language_key: str) -> List[str]:
    lang   = LANGUAGES[language_key]
    llm    = get_llm()
    parser = PydanticOutputParser(pydantic_object=SmartQuestions)
    prompt = lang["q_prompt"].format(topic=topic) + "\n\n" + parser.get_format_instructions()
    result = llm.invoke([HumanMessage(content=prompt)])
    return parser.parse(result.content).questions[:5]


def build_report_prompt(session: dict, format_instructions: str) -> str:
    topic        = session["topic"]
    lang_key     = session.get("language", "ar")
    depth        = session.get("depth", "medium")
    lang         = LANGUAGES[lang_key]
    d            = DEPTH_OPTIONS[depth]
    questions    = session.get("dynamic_questions", [])
    answers      = session.get("answers", [])
    custom_title = session.get("custom_title")

    title_instruction = (
        f'TITLE: Use EXACTLY this title: "{custom_title}" — do not change it.'
        if custom_title else "TITLE: Generate a concise academic title."
    )

    qa_block = ""
    for i, (q, a) in enumerate(zip(questions, answers), 1):
        qa_block += f"Q{i}: {q}\nA{i}: {a}\n\n"

    comparison_injection = ""
    if session.get("comparison_query"):
        cq = session["comparison_query"]
        comparison_injection = (
            f"\n\n══════════════════════════════════════\n"
            f"MANDATORY COMPARISON BLOCK — DO NOT SKIP:\n"
            f"You MUST include a 'comparison' block that compares: {cq}\n"
            f"- Set side_a and side_b to the two items\n"
            f"- Include 5-6 meaningful criteria\n"
            f"- Place this block in the middle of the report\n"
            f"══════════════════════════════════════"
        )

    # تعليمات إضافية لمنع انقسام الجداول عبر الصفحات
    table_instruction = (
        "\n\nIMPORTANT: For any table, comparison, or stats block, "
        "limit the number of rows to MAXIMUM 6 to ensure the entire table fits on one page. "
        "If you need more data, split into multiple blocks or use a different block type. "
        "NEVER let a table break across pages."
    )

    return f"""You are a skilled academic writer. Write a university report that feels GENUINELY HUMAN-WRITTEN.

══════════════════════════════════════
TOPIC: {topic}
LANGUAGE: {lang["instruction"]}
{title_instruction}
TARGET: {d["pages"]} A4 pages — approximately {d["pages"] * 400} words total.
SECTIONS: {d["blocks_min"]} to {d["blocks_max"]} content blocks.
══════════════════════════════════════

STUDENT'S REQUIREMENTS:
{qa_block.strip()}
{comparison_injection}

══════════════════════════════════════
BLOCK TYPES:
- "paragraph"     → "text": 150-200 words. Use \\n for natural breaks (3-5 times).
- "bullets"       → "items": 5-7. 40% with " — " sub-note, 60% standalone.
- "numbered_list" → "items": 5-7. Same rule.
- "table"         → "headers" + "rows" (max 6 rows). Max 2 per report.
- "pros_cons"     → "pros": 4-5, "cons": 4-5. Style A/B/C/D. NEVER on half-empty page.
- "comparison"    → side_a, side_b, criteria 5-6 (max 6 rows). Max 2 per report.
- "stats"         → "items": 5-6. "Label: value — context". Max 6 rows.
- "examples"      → "items": 5-6.
- "quote"         → "text": 2-3 sharp sentences.

{table_instruction}

PAGE FILLING (font is large — fills fast):
• After short block → NEXT must be paragraph 150-200 words.
• Never two consecutive short blocks.
• 45% paragraphs | 35% lists | 20% tables (max 2 total).

STYLE:
• No openers: "يتناول" / "In this report" FORBIDDEN.
• Paragraphs: strong claim → develop → twist.
• Conclusion: unexpected forward-looking angle. NOT a summary.

ALL text in specified language. Conclusion MANDATORY.
{format_instructions}"""


def generate_report(session: dict):
    try:
        llm    = get_llm()
        parser = PydanticOutputParser(pydantic_object=DynamicReport)
        prompt = build_report_prompt(session, parser.get_format_instructions())
        report = None
        for attempt in range(3):
            try:
                result = llm.invoke([HumanMessage(content=prompt)])
                report = parser.parse(result.content)
                break
            except Exception as e:
                if attempt == 2:
                    raise e
                logger.warning(f"Parse attempt {attempt+1} failed: {e}")
        html_str  = render_html(report, session)
        pdf_bytes = WeasyHTML(string=html_str).write_pdf()
        return pdf_bytes, report.title
    except Exception as e:
        logger.error(f"❌ generate_report: {e}", exc_info=True)
        return None, str(e)


def generate_word_report(session: dict):
    """
    توليد تقرير بصيغة Word باستخدام python-docx.
    """
    try:
        # أولاً نولد التقرير النصي باستخدام LLM (نفس الخطوات)
        llm    = get_llm()
        parser = PydanticOutputParser(pydantic_object=DynamicReport)
        prompt = build_report_prompt(session, parser.get_format_instructions())
        report = None
        for attempt in range(3):
            try:
                result = llm.invoke([HumanMessage(content=prompt)])
                report = parser.parse(result.content)
                break
            except Exception as e:
                if attempt == 2:
                    raise e
                logger.warning(f"Parse attempt {attempt+1} failed: {e}")

        # إنشاء مستند Word
        doc = Document()

        # ضبط اتجاه المستند حسب اللغة
        lang_key = session.get("language", "ar")
        if lang_key == "ar":
            # للعربية نضبط الاتجاه من اليمين لليسار
            section = doc.sections[0]
            section.right_to_left = True

        # إضافة العنوان
        title = doc.add_heading(report.title, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # إضافة المقدمة
        doc.add_heading(LANGUAGES[lang_key]['intro_label'], level=2)
        doc.add_paragraph(report.introduction)

        # إضافة الكتل
        for block in report.blocks:
            doc.add_heading(block.title, level=3)

            if block.block_type == "paragraph":
                doc.add_paragraph(block.text)

            elif block.block_type in ("bullets", "numbered_list"):
                for item in block.items or []:
                    doc.add_paragraph(item, style='List Bullet' if block.block_type == "bullets" else 'List Number')

            elif block.block_type == "table":
                if block.headers and block.rows:
                    table = doc.add_table(rows=1, cols=len(block.headers))
                    table.style = 'Light Grid Accent 1'
                    hdr_cells = table.rows[0].cells
                    for i, h in enumerate(block.headers):
                        hdr_cells[i].text = h
                    for row_data in block.rows:
                        row_cells = table.add_row().cells
                        for i, cell_text in enumerate(row_data):
                            if i < len(row_data):
                                row_cells[i].text = cell_text

            elif block.block_type == "pros_cons":
                doc.add_paragraph("المزايا:", style='List Bullet')
                for p in block.pros or []:
                    doc.add_paragraph(p, style='List Bullet')
                doc.add_paragraph("العيوب:", style='List Bullet')
                for c in block.cons or []:
                    doc.add_paragraph(c, style='List Bullet')

            elif block.block_type == "comparison":
                if block.criteria and block.side_a_values and block.side_b_values:
                    table = doc.add_table(rows=1, cols=3)
                    table.style = 'Light Grid Accent 1'
                    hdr = table.rows[0].cells
                    hdr[0].text = "المعيار"
                    hdr[1].text = block.side_a or "A"
                    hdr[2].text = block.side_b or "B"
                    for i, crit in enumerate(block.criteria):
                        row = table.add_row().cells
                        row[0].text = crit
                        row[1].text = block.side_a_values[i] if i < len(block.side_a_values) else "-"
                        row[2].text = block.side_b_values[i] if i < len(block.side_b_values) else "-"

            elif block.block_type == "stats":
                for item in block.items or []:
                    doc.add_paragraph(item, style='List Bullet')

            elif block.block_type == "examples":
                for item in block.items or []:
                    doc.add_paragraph(item, style='List Bullet')

            elif block.block_type == "quote":
                doc.add_paragraph(block.text, style='Intense Quote')

        # إضافة الخاتمة
        doc.add_heading(LANGUAGES[lang_key]['conclusion_label'], level=2)
        doc.add_paragraph(report.conclusion)

        # حفظ المستند في BytesIO
        docx_bytes = BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        return docx_bytes.getvalue(), report.title

    except Exception as e:
        logger.error(f"❌ generate_word_report: {e}", exc_info=True)
        return None, str(e)


# ═══════════════════════════════════════════════════
# HTML RENDERER
# ═══════════════════════════════════════════════════
def esc(v):
    return html_lib.escape(str(v)) if v is not None else ""

def render_item_with_subnote(item: str, txt_color: str, accent: str) -> str:
    sep = " — "
    if sep in str(item):
        parts = str(item).split(sep, 1)
        return (
            f'{esc(parts[0].strip())}'
            f'<span style="color:{accent};font-size:0.88em;font-weight:normal;"> — {esc(parts[1].strip())}</span>'
        )
    return esc(item)

def text_to_paras(text: str, align: str) -> str:
    lines = [l.strip() for l in str(text).split('\n') if l.strip()]
    if not lines:
        lines = [str(text)]
    return "".join(
        f'<p style="text-align:{align};margin:0 0 10px 0;line-height:2.05;">{esc(l)}</p>'
        for l in lines
    )

def render_block(b: ReportBlock, tc: dict, lang: dict) -> str:
    p      = tc["primary"]
    a      = tc["accent"]
    bg     = tc["bg"]
    bg2    = tc["bg2"]
    align  = lang["align"]
    is_rtl = lang["dir"] == "rtl"
    b_side = "border-right" if is_rtl else "border-left"
    p_side = "padding-right" if is_rtl else "padding-left"
    is_dark   = (p == "#d4af37")
    txt_color = "#e2e8f0" if is_dark else "#333333"
    h2_bg     = "#3d4a5c" if is_dark else bg

    h2 = (
        f'<h2 style="color:{p};font-size:15.5px;font-weight:bold;'
        f'padding:10px 16px;background:{h2_bg};'
        f'{b_side}:5px solid {a};margin:0 0 13px 0;">'
        f'{esc(b.title)}</h2>'
    )
    bt = (b.block_type or "paragraph").strip().lower()

    if bt == "paragraph":
        return f'<div style="margin:18px 0;">{h2}{text_to_paras(b.text or "", align)}</div>'

    elif bt in ("bullets", "numbered_list"):
        tag = "ol" if bt == "numbered_list" else "ul"
        lis = "".join(
            f'<li style="margin-bottom:10px;line-height:1.95;color:{txt_color};">'
            f'{render_item_with_subnote(i, txt_color, a)}</li>'
            for i in (b.items or [])
        )
        return f'<div style="margin:18px 0;">{h2}<{tag} style="{p_side}:22px;margin:0;">{lis}</{tag}></div>'

    elif bt == "stats":
        rows = ""
        for idx, item in enumerate(b.items or []):
            parts = str(item).split(":", 1)
            bg_r  = bg if idx % 2 == 0 else bg2
            if len(parts) == 2:
                rows += (
                    f'<tr><td style="font-weight:bold;color:{p};padding:9px 12px;background:{bg};'
                    f'border:1px solid #ddd;width:36%;">{esc(parts[0].strip())}</td>'
                    f'<td style="padding:9px 12px;border:1px solid #ddd;background:{bg_r};'
                    f'color:{txt_color};">{esc(parts[1].strip())}</td></tr>'
                )
            else:
                rows += f'<tr><td colspan="2" style="padding:9px 12px;border:1px solid #ddd;">{esc(item)}</td></tr>'
        return (
            f'<div class="block-stats" style="margin:18px 0;page-break-inside:avoid;">{h2}'
            f'<table style="width:100%;border-collapse:collapse;font-size:14px;">{rows}</table></div>'
        )

    elif bt == "examples":
        rows = ""
        for idx, item in enumerate(b.items or [], 1):
            bg_r = bg if idx % 2 == 1 else bg2
            rows += (
                f'<tr><td style="width:28px;text-align:center;font-weight:bold;color:#fff;'
                f'background:{a};padding:9px;border:1px solid #ddd;">{idx}</td>'
                f'<td style="padding:9px 12px;border:1px solid #ddd;background:{bg_r};'
                f'line-height:1.95;color:{txt_color};">{render_item_with_subnote(item, txt_color, a)}</td></tr>'
            )
        return (
            f'<div style="margin:18px 0;">{h2}'
            f'<table style="width:100%;border-collapse:collapse;font-size:14px;">{rows}</table></div>'
        )

    elif bt == "pros_cons":
        pros  = b.pros or []
        cons  = b.cons or []
        style = (b.style or "A").upper().strip()

        def pro_li(x):
            sep = " — "
            if sep in str(x):
                pts = str(x).split(sep, 1)
                return (
                    f'<li style="margin-bottom:8px;line-height:1.85;font-size:14px;">'
                    f'<span style="font-weight:700;color:#1a5e38;">{esc(pts[0].strip())}</span>'
                    f'<br><span style="color:#2d6a4f;font-size:13px;{p_side}:6px;">↳ {esc(pts[1].strip())}</span></li>'
                )
            return f'<li style="margin-bottom:8px;line-height:1.85;font-size:14px;font-weight:600;color:#1a5e38;">{esc(x)}</li>'

        def con_li(x):
            sep = " — "
            if sep in str(x):
                pts = str(x).split(sep, 1)
                return (
                    f'<li style="margin-bottom:8px;line-height:1.85;font-size:14px;">'
                    f'<span style="font-weight:700;color:#7b1a1a;">{esc(pts[0].strip())}</span>'
                    f'<br><span style="color:#922b21;font-size:13px;{p_side}:6px;">↳ {esc(pts[1].strip())}</span></li>'
                )
            return f'<li style="margin-bottom:8px;line-height:1.85;font-size:14px;font-weight:600;color:#7b1a1a;">{esc(x)}</li>'

        if style == "A":
            p_lis = "".join(pro_li(x) for x in pros)
            c_lis = "".join(con_li(x) for x in cons)
            inner = (
                f'<table style="width:100%;border-collapse:separate;border-spacing:8px 0;"><tr>'
                f'<td style="vertical-align:top;width:50%;padding:0;">'
                f'<div style="background:#1a5e38;color:#fff;font-weight:700;font-size:14px;padding:9px 16px;border-radius:6px 6px 0 0;">{lang["pros_label"]}</div>'
                f'<div style="background:#f0fff4;border:2px solid #1a5e38;border-top:none;border-radius:0 0 6px 6px;padding:10px 14px;">'
                f'<ul style="{p_side}:14px;margin:0;">{p_lis}</ul></div></td>'
                f'<td style="vertical-align:top;width:50%;padding:0;">'
                f'<div style="background:#7b1a1a;color:#fff;font-weight:700;font-size:14px;padding:9px 16px;border-radius:6px 6px 0 0;">{lang["cons_label"]}</div>'
                f'<div style="background:#fff5f5;border:2px solid #7b1a1a;border-top:none;border-radius:0 0 6px 6px;padding:10px 14px;">'
                f'<ul style="{p_side}:14px;margin:0;">{c_lis}</ul></div></td>'
                f'</tr></table>'
            )
        elif style == "B":
            rows_html = ""
            for sign, item in [("+", x) for x in pros] + [("-", x) for x in cons]:
                is_pro   = sign == "+"
                row_bg   = "#f0fff4" if is_pro else "#fff5f5"
                dot_bg   = "#1a5e38" if is_pro else "#7b1a1a"
                dot_char = "✓" if is_pro else "✗"
                sep = " — "
                if sep in str(item):
                    pts  = str(item).split(sep, 1)
                    cell = f'<span style="font-weight:700;">{esc(pts[0].strip())}</span><span style="color:#555;font-size:13px;"> — {esc(pts[1].strip())}</span>'
                else:
                    cell = f'<span style="font-weight:600;">{esc(item)}</span>'
                rows_html += (
                    f'<tr style="background:{row_bg};">'
                    f'<td style="width:32px;text-align:center;font-weight:800;color:{dot_bg};font-size:17px;padding:10px 6px;border-bottom:1px solid #e8e8e8;">{dot_char}</td>'
                    f'<td style="padding:10px 12px;border-bottom:1px solid #e8e8e8;font-size:14px;line-height:1.8;">{cell}</td></tr>'
                )
            inner = (
                f'<table style="width:100%;border-collapse:collapse;border:1px solid #d0d0d0;">'
                f'<thead><tr>'
                f'<th style="background:#2d3748;color:#fff;padding:9px 6px;width:32px;font-size:14px;">±</th>'
                f'<th style="background:#2d3748;color:#fff;padding:9px 14px;text-align:{align};font-size:14px;">التفاصيل</th>'
                f'</tr></thead><tbody>{rows_html}</tbody></table>'
            )
        elif style == "C":
            p_lis = "".join(pro_li(x) for x in pros)
            c_lis = "".join(con_li(x) for x in cons)
            inner = (
                f'<div style="border:2px solid #1a5e38;border-radius:8px;margin-bottom:10px;">'
                f'<div style="background:#1a5e38;color:#fff;font-weight:700;font-size:14px;padding:9px 16px;border-radius:6px 6px 0 0;">{lang["pros_label"]}</div>'
                f'<div style="background:#f0fff4;padding:10px 16px;"><ul style="{p_side}:16px;margin:0;">{p_lis}</ul></div></div>'
                f'<div style="border:2px solid #7b1a1a;border-radius:8px;">'
                f'<div style="background:#7b1a1a;color:#fff;font-weight:700;font-size:14px;padding:9px 16px;border-radius:6px 6px 0 0;">{lang["cons_label"]}</div>'
                f'<div style="background:#fff5f5;padding:10px 16px;"><ul style="{p_side}:16px;margin:0;">{c_lis}</ul></div></div>'
            )
        else:  # D
            items_html = ""
            for emoji, lst in [("✅", pros), ("❌", cons)]:
                for x in lst:
                    sep = " — "
                    if sep in str(x):
                        pts = str(x).split(sep, 1)
                        t   = f'<b>{esc(pts[0].strip())}</b> — <span style="color:#555;">{esc(pts[1].strip())}</span>'
                    else:
                        t = f'<b>{esc(x)}</b>'
                    items_html += (
                        f'<div style="display:flex;gap:10px;margin-bottom:10px;align-items:flex-start;">'
                        f'<span style="font-size:17px;flex-shrink:0;">{emoji}</span>'
                        f'<span style="font-size:14px;line-height:1.85;">{t}</span></div>'
                    )
            inner = f'<div style="background:{bg};{b_side}:3px solid {a};padding:14px 18px;border-radius:6px;">{items_html}</div>'

        return f'<div style="margin:18px 0;">{h2}{inner}</div>'

    elif bt == "table":
        ths = "".join(
            f'<th style="background:{p};color:#fff;padding:10px 12px;text-align:{align};font-weight:bold;">{esc(h)}</th>'
            for h in (b.headers or [])
        )
        rows = ""
        for ridx, row in enumerate(b.rows or []):
            bg_r = bg if ridx % 2 == 0 else bg2
            tds  = "".join(
                f'<td style="padding:9px 12px;border:1px solid #ddd;background:{bg_r};color:{txt_color};">{esc(c)}</td>'
                for c in row
            )
            rows += f"<tr>{tds}</tr>"
        return (
            f'<div class="block-table" style="margin:18px 0;page-break-inside:avoid;">{h2}'
            f'<table style="width:100%;border-collapse:collapse;font-size:14px;">'
            f'<thead><tr>{ths}</tr></thead><tbody>{rows}</tbody></table></div>'
        )

    elif bt == "comparison":
        sa = esc(b.side_a or "A")
        sb = esc(b.side_b or "B")
        cr = b.criteria or []
        av = b.side_a_values or []
        bv = b.side_b_values or []
        ths = (
            f'<th style="background:{p};color:#fff;padding:10px 12px;">المعيار</th>'
            f'<th style="background:{p};color:#fff;padding:10px 12px;">{sa}</th>'
            f'<th style="background:{p};color:#fff;padding:10px 12px;">{sb}</th>'
        )
        rows = ""
        for idx, crit in enumerate(cr):
            bg_r = bg if idx % 2 == 0 else bg2
            rows += (
                f'<tr><td style="font-weight:bold;color:{p};padding:9px 12px;border:1px solid #ddd;background:{bg};">{esc(crit)}</td>'
                f'<td style="padding:9px 12px;border:1px solid #ddd;background:{bg_r};">{esc(av[idx]) if idx < len(av) else "-"}</td>'
                f'<td style="padding:9px 12px;border:1px solid #ddd;background:{bg_r};">{esc(bv[idx]) if idx < len(bv) else "-"}</td></tr>'
            )
        return (
            f'<div class="block-comparison" style="margin:18px 0;page-break-inside:avoid;">{h2}'
            f'<table style="width:100%;border-collapse:collapse;font-size:14px;">'
            f'<thead><tr>{ths}</tr></thead><tbody>{rows}</tbody></table></div>'
        )

    elif bt == "quote":
        bd = "border-right" if is_rtl else "border-left"
        pd = "padding-right" if is_rtl else "padding-left"
        return (
            f'<div style="margin:18px 0;">{h2}'
            f'<blockquote style="{bd}:5px solid {a};{pd}:16px;margin:0;'
            f'color:#555;font-style:italic;line-height:2.0;">{esc(b.text or "")}</blockquote></div>'
        )

    else:
        return f'<div style="margin:18px 0;">{h2}{text_to_paras(b.text or "", align)}</div>'


# ═══════════════════════════════════════════════════
# RENDER HTML
# ═══════════════════════════════════════════════════
def render_html(report: DynamicReport, session: dict) -> str:
    language_key  = session.get("language", "ar")
    lang          = LANGUAGES[language_key]
    is_custom     = session.get("custom_mode", False)
    template_name = "_custom" if is_custom else session.get("template", "emerald")

    if is_custom:
        colors    = CUSTOM_COLORS[session.get("custom_color_key", "royal_blue")]
        p, a, bg, bg2 = colors["primary"], colors["accent"], colors["bg"], colors["bg2"]
        font_size = CUSTOM_FONT_SIZES[session.get("custom_font_size_key", "medium")]["size"]
        font      = CUSTOM_FONTS[session.get("custom_font_key", "traditional")]["value"]
        line_height = LINE_HEIGHTS[session.get("custom_line_height", "normal")]["value"]
        page_margin = PAGE_MARGINS[session.get("custom_page_margin", "medium")]["value"]
        header_style_key = session.get("custom_header_style", "colored")
        header_color = HEADER_STYLES[header_style_key]["color"]
        if header_color == "auto":
            header_color = a
        header_size = HEADER_STYLES[header_style_key]["size"]
        show_hf = SHOW_HEADER_FOOTER[session.get("custom_show_header_footer", "yes")]["show"]
    else:
        tc        = TEMPLATES[template_name]
        p, a, bg, bg2 = tc["primary"], tc["accent"], tc["bg"], tc["bg2"]
        font_size = "16.5px"
        font      = lang["font"]
        line_height = "1.8"  # normal
        page_margin = "2.5cm" # medium
        header_color = p
        header_size = "24px"
        show_hf = True

    tc     = {"primary": p, "accent": a, "bg": bg, "bg2": bg2}
    dir_   = lang["dir"]
    align  = lang["align"]
    is_rtl = dir_ == "rtl"
    b_side = "border-right" if is_rtl else "border-left"

    # ── ألوان الصفحة ──
    if template_name == "dark_elegant":
        page_bg, body_color, box_bg = "#1a202c", "#e2e8f0", "#2d3748"
    elif template_name == "royal":
        page_bg, body_color, box_bg = "#fffdf7", "#2c1810", "#fdf6e3"
    else:
        page_bg, body_color, box_bg = "#ffffff", "#2d3436", bg

    # ── إطار الصفحة ──
    borders = {
        "emerald":      (f"3px solid {p}", "0.35cm", "0.7cm", f"outline:1.5px solid {a};outline-offset:-7px;"),
        "modern":       (f"4px solid {a}", "0.35cm", "0.7cm", ""),
        "minimal":      (f"1.5px solid {p}","0.4cm",  "0.7cm", ""),
        "professional": (f"2px solid {p}", "0.35cm", "0.65cm",f"outline:4px solid {p};outline-offset:-10px;"),
        "dark_elegant": (f"2px solid {a}", "0.35cm", "0.7cm", ""),
        "royal":        (f"3px solid {p}", "0.35cm", "0.7cm", f"outline:2px solid {a};outline-offset:-8px;"),
        "_custom":      (f"3px solid {p}", "0.35cm", "0.7cm", f"outline:1.5px solid {a};outline-offset:-8px;"),
    }
    page_border, page_margin_extra, page_padding, extra_css = borders.get(
        template_name, ("none", "2cm", "0cm", "")
    )
    # استخدم page_margin المخصص إذا كان موجوداً
    final_margin = page_margin if is_custom else page_margin_extra

    # ── ترويسة وتذييل ──
    gdir = "left" if is_rtl else "right"
    if template_name == "professional":
        prof_top = (
            f'<div style="margin-bottom:24px;">'
            f'<div style="height:5px;background:{p};"></div>'
            f'<div style="height:2px;background:{a};margin-top:3px;"></div>'
            f'<div style="display:flex;justify-content:space-between;align-items:center;padding:8px 4px 6px 4px;">'
            f'<span style="font-size:11px;color:{a};font-weight:700;letter-spacing:2px;">تقرير أكاديمي رسمي</span>'
            f'<span style="font-size:10px;color:#8b9bb4;letter-spacing:1px;">OFFICIAL ACADEMIC REPORT</span>'
            f'</div><div style="height:1px;background:#d0dae8;"></div></div>'
        )
        prof_bot = (
            f'<div style="margin-top:24px;">'
            f'<div style="height:1px;background:#d0dae8;"></div>'
            f'<div style="display:flex;justify-content:space-between;padding:6px 4px;">'
            f'<span style="font-size:10px;color:#8b9bb4;">سري — للاستخدام الأكاديمي فقط</span>'
            f'<span style="font-size:10px;color:#8b9bb4;">Confidential — Academic Use Only</span>'
            f'</div>'
            f'<div style="height:2px;background:{a};"></div>'
            f'<div style="height:5px;background:{p};margin-top:3px;"></div></div>'
        )
    elif template_name == "royal":
        prof_top = (
            f'<div style="margin-bottom:22px;text-align:center;">'
            f'<div style="height:4px;background:linear-gradient(to {gdir},{p},{a},{p});border-radius:2px;"></div>'
            f'<div style="padding:8px 4px 5px;"><span style="font-size:12px;color:{a};font-weight:700;letter-spacing:3px;">✦ تقرير أكاديمي جامعي ✦</span></div>'
            f'<div style="height:1px;background:{a};opacity:0.35;"></div></div>'
        )
        prof_bot = (
            f'<div style="margin-top:22px;text-align:center;">'
            f'<div style="height:1px;background:{a};opacity:0.35;"></div>'
            f'<div style="padding:6px 4px;"><span style="font-size:11px;color:{a};letter-spacing:2px;">✦ إعداد أكاديمي رسمي — جميع الحقوق محفوظة ✦</span></div>'
            f'<div style="height:4px;background:linear-gradient(to {gdir},{p},{a},{p});border-radius:2px;"></div></div>'
        )
    elif template_name == "_custom":
        prof_top = (
            f'<div style="margin-bottom:16px;">'
            f'<div style="height:3px;background:linear-gradient(to {gdir},{p},{a},{p});border-radius:2px;"></div>'
            f'</div>'
        )
        prof_bot = (
            f'<div style="margin-top:16px;">'
            f'<div style="height:3px;background:linear-gradient(to {gdir},{p},{a},{p});border-radius:2px;"></div>'
            f'</div>'
        )
    else:
        prof_top = prof_bot = ""

    if not show_hf:
        prof_top = ""
        prof_bot = ""

    blocks_html = "\n".join(render_block(bl, tc, lang) for bl in report.blocks)

    return f"""<!DOCTYPE html>
<html lang="{lang['lang_attr']}" dir="{dir_}">
<head>
<meta charset="UTF-8">
<style>
  @page {{
    size: A4;
    margin: {final_margin};
    border: {page_border};
    padding: {page_padding};
    background: {page_bg};
    {extra_css}
  }}
  * {{ box-sizing: border-box; }}
  body {{
    font-family: {font};
    direction: {dir_};
    text-align: justify;
    line-height: {line_height};
    color: {body_color};
    background: {page_bg};
    font-size: {font_size};
    margin: 0; padding: 0;
    word-spacing: 0.05em;
  }}
  p  {{ text-align: justify; margin: 0 0 9px 0; }}
  h1 {{ font-size: {header_size} !important; text-align: center; color: {header_color}; }}
  h2 {{ font-size: 15.5px !important; text-align: {align}; }}
  li {{ text-align: {align}; }}
  p, li {{ orphans: 2; widows: 2; }}
  .block-table, .block-stats, .block-comparison {{ page-break-inside: avoid; }}
  h2 {{ page-break-after: avoid; orphans: 3; widows: 3; }}
</style>
</head>
<body>

{prof_top}

<h1 style="text-align:center; padding-bottom:14px; margin-bottom:28px; border-bottom:3px solid {a};">
  {esc(report.title)}
</h1>

<div style="background:{box_bg};padding:18px 22px;border-radius:8px;
            margin:0 0 20px 0;{b_side}:5px solid {a};">
  <h2 style="color:{p};font-size:15.5px;font-weight:bold;margin:0 0 10px 0;">
    📚 {lang['intro_label']}
  </h2>
  {text_to_paras(report.introduction, align)}
</div>

{blocks_html}

<div style="background:{box_bg};padding:18px 22px;border-radius:8px;
            margin:20px 0 0 0;{b_side}:5px solid {a};">
  <h2 style="color:{p};font-size:15.5px;font-weight:bold;margin:0 0 10px 0;">
    🎯 {lang['conclusion_label']}
  </h2>
  {text_to_paras(report.conclusion, align)}
</div>

{prof_bot}

</body>
</html>"""


# ═══════════════════════════════════════════════════
# KEYBOARDS
# ═══════════════════════════════════════════════════
def title_keyboard():
    return InlineKeyboardMarkup([[InlineKeyboardButton("👻 اتركه للشبح", callback_data="title_auto")]])

def lang_keyboard():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton(v["name"], callback_data=f"lang_{k}")]
        for k, v in LANGUAGES.items()
    ])

def depth_keyboard():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton(v["name"], callback_data=f"depth_{k}")]
        for k, v in DEPTH_OPTIONS.items()
    ])

def style_mode_keyboard():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("🎭 قوالب جاهزة",   callback_data="style_preset")],
        [InlineKeyboardButton("🎨 تخصيص كامل ✨", callback_data="style_custom")],
    ])

def template_keyboard():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton(v["name"], callback_data=f"tpl_{k}")]
        for k, v in TEMPLATES.items()
    ])

def font_size_keyboard():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton(v["label"], callback_data=f"fsize_{k}")]
        for k, v in CUSTOM_FONT_SIZES.items()
    ])

def font_keyboard():
    items = list(CUSTOM_FONTS.items())
    rows  = []
    for i in range(0, len(items), 2):
        row = [InlineKeyboardButton(items[i][1]["label"], callback_data=f"cfont_{items[i][0]}")]
        if i + 1 < len(items):
            row.append(InlineKeyboardButton(items[i+1][1]["label"], callback_data=f"cfont_{items[i+1][0]}"))
        rows.append(row)
    return InlineKeyboardMarkup(rows)

def colors_keyboard():
    items = list(CUSTOM_COLORS.items())
    rows  = []
    for i in range(0, len(items), 2):
        row = [InlineKeyboardButton(items[i][1]["label"], callback_data=f"color_{items[i][0]}")]
        if i + 1 < len(items):
            row.append(InlineKeyboardButton(items[i+1][1]["label"], callback_data=f"color_{items[i+1][0]}"))
        rows.append(row)
    return InlineKeyboardMarkup(rows)

def line_height_keyboard():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton(v["label"], callback_data=f"lh_{k}")]
        for k, v in LINE_HEIGHTS.items()
    ])

def page_margin_keyboard():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton(v["label"], callback_data=f"pm_{k}")]
        for k, v in PAGE_MARGINS.items()
    ])

def header_style_keyboard():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton(v["label"], callback_data=f"hs_{k}")]
        for k, v in HEADER_STYLES.items()
    ])

def show_header_keyboard():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton(v["label"], callback_data=f"sh_{k}")]
        for k, v in SHOW_HEADER_FOOTER.items()
    ])

def file_format_keyboard():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton(f"{FILE_FORMATS['pdf']['emoji']} PDF", callback_data="format_pdf"),
         InlineKeyboardButton(f"{FILE_FORMATS['word']['emoji']} Word", callback_data="format_word")]
    ])

def comparison_keyboard():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("📊 نعم، أضف جدول مقارنة!", callback_data="comp_yes")],
        [InlineKeyboardButton("❌ لا شكراً",               callback_data="comp_no")],
    ])


# ═══════════════════════════════════════════════════
# HELPER
# ═══════════════════════════════════════════════════
def build_queue_text(session: dict, pos: int) -> str:
    lang_name  = LANGUAGES[session.get("language", "ar")]["name"]
    depth_name = DEPTH_OPTIONS[session.get("depth", "medium")]["name"]
    tpl_name   = "🎨 مخصص" if session.get("custom_mode") else TEMPLATES.get(session.get("template", "emerald"), {}).get("name", "")
    safe_topic = session["topic"].replace('<','&lt;').replace('>','&gt;').replace('&','&amp;')
    file_format = FILE_FORMATS[session.get("file_format", "pdf")]["emoji"]
    status     = "👻 <b>الشبح بدأ يكتب تقريرك!</b>" if pos == 1 else f"⏳ <b>في الطابور — الترتيب {pos}</b> 👻"
    return f"{status}\n\n📝 <b>الموضوع:</b> <i>{safe_topic}</i>\n🌐 {lang_name}  |  📏 {depth_name}  |  🎨 {tpl_name}  |  {file_format}"


# ═══════════════════════════════════════════════════
# TELEGRAM HANDLERS
# ═══════════════════════════════════════════════════
async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    if user_id in user_sessions:
        user_sessions.pop(user_id, None)
        queue_positions.pop(user_id, None)
        await update.message.reply_text("❌ <b>تم إلغاء الجلسة.</b>\n\n👻 أرسل موضوعاً جديداً لبدء تقرير جديد.", parse_mode='HTML')
    else:
        await update.message.reply_text("ℹ️ لا توجد جلسة نشطة.\n\n👻 أرسل موضوعاً لبدء تقرير جديد.", parse_mode='HTML')


async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    user_sessions.pop(user_id, None)
    name = update.effective_user.first_name
    await update.message.reply_text(
        f"👻 <b>أهلاً {name}! أنا Repooreto</b>\n"
        f"الشبح الذي يكتب تقاريرك الجامعية! 🎓\n\n"
        f"✨ <b>كيف أعمل؟</b>\n"
        f"1️⃣ أرسل موضوع تقريرك\n"
        f"2️⃣ اختر اللغة\n"
        f"3️⃣ أجب على أسئلتي الذكية 🧠\n"
        f"4️⃣ اختر العمق والتصميم:\n"
        f"   • 🎭 <b>قوالب جاهزة</b> — 6 قوالب احترافية\n"
        f"   • 🎨 <b>تخصيص كامل</b> — خط، ألوان، مقارنة خاصة ✨\n"
        f"5️⃣ اختر صيغة الملف: PDF أو Word\n"
        f"6️⃣ استلم تقريرك! 🎉\n\n"
        f"👻 <b>أرسل موضوع تقريرك الآن!</b>",
        parse_mode='HTML'
    )


async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    text    = update.message.text.strip()

    if user_id in user_sessions:
        session = user_sessions[user_id]
        state   = session.get("state", "")

        if state == "answering":
            answers   = session.setdefault("answers", [])
            questions = session.get("dynamic_questions", [])
            answers.append(text)
            if len(answers) < len(questions):
                nq    = questions[len(answers)]
                q_num = len(answers) + 1
                total = len(questions)
                await update.message.reply_text(
                    f"✅ تم تسجيل إجابتك.\n\n❓ <b>السؤال {q_num}/{total}:</b>\n{nq}\n\n<i>اكتب إجابتك 👇</i>",
                    parse_mode='HTML'
                )
            else:
                session["state"] = "choosing_title"
                await update.message.reply_text(
                    "✅ <b>ممتاز! تم تسجيل جميع إجاباتك.</b>\n\n"
                    "📌 <b>هل تريد تحديد عنوان للتقرير؟</b>\n"
                    "<i>اكتب العنوان، أو دع الشبح يختاره 👇</i>",
                    reply_markup=title_keyboard(), parse_mode='HTML'
                )
            return

        if state == "choosing_title":
            session["custom_title"] = text
            session["state"]        = "choosing_depth"
            await update.message.reply_text(
                f"✅ <b>العنوان:</b> <i>{esc(text)}</i>\n\n📏 <b>اختر عمق التقرير:</b>",
                reply_markup=depth_keyboard(), parse_mode='HTML'
            )
            return

        if state == "entering_comparison":
            session["comparison_query"] = text
            session["state"]            = "choosing_file_format"
            await update.message.reply_text(
                "📁 <b>الآن اختر صيغة الملف الذي تريده:</b>",
                reply_markup=file_format_keyboard(), parse_mode='HTML'
            )
            return

        guidance = STATE_GUIDANCE.get(state, "⏳ جاري المعالجة... أرسل /cancel للبدء من جديد.")
        await update.message.reply_text(guidance, parse_mode='HTML')
        return

    if len(text) < 5:
        await update.message.reply_text("👻 الموضوع قصير جداً! أرسل موضوعاً أوضح.")
        return
    if len(text) > 250:
        await update.message.reply_text("👻 الموضوع طويل جداً! اختصره لأقل من 250 حرف.")
        return

    user_sessions[user_id] = {"topic": text, "state": "choosing_lang"}
    safe = text.replace('<','&lt;').replace('>','&gt;').replace('&','&amp;')
    await update.message.reply_text(
        f"📝 <b>الموضوع:</b> <i>{safe}</i>\n\n🌐 <b>اختر لغة التقرير:</b>",
        reply_markup=lang_keyboard(), parse_mode='HTML'
    )


async def title_auto_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query   = update.callback_query
    await query.answer()
    user_id = query.from_user.id
    if user_id not in user_sessions:
        await query.edit_message_text("❌ الجلسة منتهية."); return
    session = user_sessions[user_id]
    if session.get("state") != "choosing_title":
        await query.answer("هذا الزر لم يعد فعالاً.", show_alert=True); return
    session.pop("custom_title", None)
    session["state"] = "choosing_depth"
    await query.edit_message_text(
        "👻 <b>سيختار الشبح العنوان المناسب!</b>\n\n📏 <b>اختر عمق التقرير:</b>",
        reply_markup=depth_keyboard(), parse_mode='HTML'
    )


async def language_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query   = update.callback_query
    await query.answer()
    user_id = query.from_user.id
    lang    = query.data.replace("lang_", "")
    if user_id not in user_sessions:
        await query.edit_message_text("❌ الجلسة منتهية."); return
    session             = user_sessions[user_id]
    session["language"] = lang
    session["state"]    = "generating_questions"
    await query.edit_message_text(
        f"✅ <b>اللغة:</b> {LANGUAGES[lang]['name']}\n\n👻 <i>الشبح يحلل موضوعك ويولّد الأسئلة...</i>",
        parse_mode='HTML'
    )
    try:
        loop      = asyncio.get_event_loop()
        questions = await loop.run_in_executor(None, generate_dynamic_questions, session["topic"], lang)
        if not questions:
            raise ValueError("no questions")
        session["dynamic_questions"] = questions
        session["state"]             = "answering"
        total  = len(questions)
        q_word = "سؤال" if total == 1 else "أسئلة"
        hint   = "\n\n💡 <i>يمكنك طلب جداول، مزايا/عيوب، أو مقارنات في إجاباتك.</i>"
        await query.edit_message_text(
            f"🧠 <b>لديّ {total} {q_word} قبل الكتابة!</b>{hint}\n\n"
            f"❓ <b>السؤال 1/{total}:</b>\n{questions[0]}\n\n<i>اكتب إجابتك 👇</i>",
            parse_mode='HTML'
        )
    except Exception as e:
        logger.error(f"Questions failed: {e}", exc_info=True)
        session["dynamic_questions"] = []
        session["answers"]           = []
        session["state"]             = "choosing_depth"
        await query.edit_message_text(
            "⚠️ تعذّر توليد الأسئلة. سنكمل مباشرةً.\n\n📏 <b>اختر عمق التقرير:</b>",
            reply_markup=depth_keyboard(), parse_mode='HTML'
        )


async def depth_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query   = update.callback_query
    await query.answer()
    user_id = query.from_user.id
    depth   = query.data.replace("depth_", "")
    if user_id not in user_sessions:
        await query.edit_message_text("❌ الجلسة منتهية."); return
    if user_sessions[user_id].get("state") != "choosing_depth":
        await query.answer("هذا الزر لم يعد فعالاً.", show_alert=True); return
    user_sessions[user_id]["depth"] = depth
    user_sessions[user_id]["state"] = "choosing_style_mode"
    await query.edit_message_text(
        f"✅ <b>العمق:</b> {DEPTH_OPTIONS[depth]['name']}\n\n"
        "🎨 <b>كيف تريد تصميم تقريرك؟</b>\n\n"
        "🎭 <b>قوالب جاهزة</b> — 6 قوالب احترافية جاهزة للاستخدام\n"
        "✨ <b>تخصيص كامل</b> — اختر الخط، الألوان، وأضف مقارنة خاصة",
        reply_markup=style_mode_keyboard(), parse_mode='HTML'
    )


async def style_mode_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query   = update.callback_query
    await query.answer()
    user_id = query.from_user.id
    mode    = query.data.replace("style_", "")
    if user_id not in user_sessions:
        await query.edit_message_text("❌ الجلسة منتهية."); return
    if user_sessions[user_id].get("state") != "choosing_style_mode":
        await query.answer("هذا الزر لم يعد فعالاً.", show_alert=True); return
    session = user_sessions[user_id]
    if mode == "preset":
        session["custom_mode"] = False
        session["state"]       = "choosing_template"
        await query.edit_message_text(
            "🎭 <b>اختر قالباً من مجموعة Repooreto:</b>",
            reply_markup=template_keyboard(), parse_mode='HTML'
        )
    else:
        session["custom_mode"] = True
        # تهيئة القيم الافتراضية للتخصيص
        session["custom_font_size_key"] = "medium"
        session["custom_font_key"] = "traditional"
        session["custom_color_key"] = "royal_blue"
        session["custom_line_height"] = "normal"
        session["custom_page_margin"] = "medium"
        session["custom_header_style"] = "colored"
        session["custom_show_header_footer"] = "yes"
        session["state"] = "choosing_font_size"
        await query.edit_message_text(
            "🎨 <b>رحلة التخصيص بدأت! 👻</b>\n\n"
            "📐 <b>الخطوة 1 من 6 — حجم الخط:</b>\n"
            "اختر الحجم الذي يريح عينيك 👇",
            reply_markup=font_size_keyboard(), parse_mode='HTML'
        )


async def font_size_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query   = update.callback_query
    await query.answer()
    user_id = query.from_user.id
    key     = query.data.replace("fsize_", "")
    if user_id not in user_sessions:
        await query.edit_message_text("❌ الجلسة منتهية."); return
    if user_sessions[user_id].get("state") != "choosing_font_size":
        await query.answer("هذا الزر لم يعد فعالاً.", show_alert=True); return
    session = user_sessions[user_id]
    session["custom_font_size_key"] = key
    session["state"]                = "choosing_font"
    await query.edit_message_text(
        f"✅ <b>الحجم:</b> {CUSTOM_FONT_SIZES[key]['label']}\n\n"
        "✍️ <b>الخطوة 2 من 6 — نوع الخط:</b>\n"
        "الخطوط العلوية للعربية — السفلية للإنجليزية 👇",
        reply_markup=font_keyboard(), parse_mode='HTML'
    )


async def font_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query   = update.callback_query
    await query.answer()
    user_id = query.from_user.id
    key     = query.data.replace("cfont_", "")
    if user_id not in user_sessions:
        await query.edit_message_text("❌ الجلسة منتهية."); return
    if user_sessions[user_id].get("state") != "choosing_font":
        await query.answer("هذا الزر لم يعد فعالاً.", show_alert=True); return
    session = user_sessions[user_id]
    session["custom_font_key"] = key
    session["state"]           = "choosing_colors"
    await query.edit_message_text(
        f"✅ <b>الخط:</b> {CUSTOM_FONTS[key]['label']}\n\n"
        "🎨 <b>الخطوة 3 من 6 — نظام الألوان:</b>\n"
        "اختر الروح البصرية لتقريرك 👇",
        reply_markup=colors_keyboard(), parse_mode='HTML'
    )


async def colors_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query   = update.callback_query
    await query.answer()
    user_id = query.from_user.id
    key     = query.data.replace("color_", "")
    if user_id not in user_sessions:
        await query.edit_message_text("❌ الجلسة منتهية."); return
    if user_sessions[user_id].get("state") != "choosing_colors":
        await query.answer("هذا الزر لم يعد فعالاً.", show_alert=True); return
    session = user_sessions[user_id]
    session["custom_color_key"] = key
    session["state"]            = "choosing_line_height"
    await query.edit_message_text(
        f"✅ <b>الألوان:</b> {CUSTOM_COLORS[key]['label']}\n\n"
        "📏 <b>الخطوة 4 من 6 — تباعد الأسطر:</b>\n"
        "اختر المسافة بين السطور 👇",
        reply_markup=line_height_keyboard(), parse_mode='HTML'
    )


async def line_height_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query   = update.callback_query
    await query.answer()
    user_id = query.from_user.id
    key     = query.data.replace("lh_", "")
    if user_id not in user_sessions:
        await query.edit_message_text("❌ الجلسة منتهية."); return
    if user_sessions[user_id].get("state") != "choosing_line_height":
        await query.answer("هذا الزر لم يعد فعالاً.", show_alert=True); return
    session = user_sessions[user_id]
    session["custom_line_height"] = key
    session["state"]            = "choosing_page_margin"
    await query.edit_message_text(
        f"✅ <b>تباعد الأسطر:</b> {LINE_HEIGHTS[key]['label']}\n\n"
        "📐 <b>الخطوة 5 من 6 — هوامش الصفحة:</b>\n"
        "اختر حجم الهوامش 👇",
        reply_markup=page_margin_keyboard(), parse_mode='HTML'
    )


async def page_margin_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query   = update.callback_query
    await query.answer()
    user_id = query.from_user.id
    key     = query.data.replace("pm_", "")
    if user_id not in user_sessions:
        await query.edit_message_text("❌ الجلسة منتهية."); return
    if user_sessions[user_id].get("state") != "choosing_page_margin":
        await query.answer("هذا الزر لم يعد فعالاً.", show_alert=True); return
    session = user_sessions[user_id]
    session["custom_page_margin"] = key
    session["state"]            = "choosing_header_style"
    await query.edit_message_text(
        f"✅ <b>الهوامش:</b> {PAGE_MARGINS[key]['label']}\n\n"
        "🎯 <b>الخطوة 6 من 6 — نمط العنوان الرئيسي:</b>\n"
        "اختر شكل العنوان 👇",
        reply_markup=header_style_keyboard(), parse_mode='HTML'
    )


async def header_style_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query   = update.callback_query
    await query.answer()
    user_id = query.from_user.id
    key     = query.data.replace("hs_", "")
    if user_id not in user_sessions:
        await query.edit_message_text("❌ الجلسة منتهية."); return
    if user_sessions[user_id].get("state") != "choosing_header_style":
        await query.answer("هذا الزر لم يعد فعالاً.", show_alert=True); return
    session = user_sessions[user_id]
    session["custom_header_style"] = key
    session["state"]            = "choosing_show_header"
    await query.edit_message_text(
        f"✅ <b>نمط العنوان:</b> {HEADER_STYLES[key]['label']}\n\n"
        "📰 <b>هل تريد إظهار الترويسة والتذييل؟</b>",
        reply_markup=show_header_keyboard(), parse_mode='HTML'
    )


async def show_header_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query   = update.callback_query
    await query.answer()
    user_id = query.from_user.id
    key     = query.data.replace("sh_", "")
    if user_id not in user_sessions:
        await query.edit_message_text("❌ الجلسة منتهية."); return
    if user_sessions[user_id].get("state") != "choosing_show_header":
        await query.answer("هذا الزر لم يعد فعالاً.", show_alert=True); return
    session = user_sessions[user_id]
    session["custom_show_header_footer"] = key
    session["state"]            = "asking_comparison"
    await query.edit_message_text(
        f"✅ <b>الترويسة:</b> {SHOW_HEADER_FOOTER[key]['label']}\n\n"
        "📊 <b>هل تريد إضافة جدول مقارنة خاص في التقرير؟</b>\n"
        "<i>مثال: مقارنة Python مع Java، أو الطاقة الشمسية مع النووية...</i>",
        reply_markup=comparison_keyboard(), parse_mode='HTML'
    )


async def comp_yes_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query   = update.callback_query
    await query.answer()
    user_id = query.from_user.id
    if user_id not in user_sessions:
        await query.edit_message_text("❌ الجلسة منتهية."); return
    if user_sessions[user_id].get("state") != "asking_comparison":
        await query.answer("هذا الزر لم يعد فعالاً.", show_alert=True); return
    user_sessions[user_id]["state"] = "entering_comparison"
    await query.edit_message_text(
        "📊 <b>اكتب الشيئين اللذين تريد مقارنتهما:</b>\n\n"
        "💡 <i>أمثلة:</i>\n"
        "• <code>Python مقابل Java</code>\n"
        "• <code>العمل الحر مقابل الوظيفة</code>\n"
        "• <code>الطاقة الشمسية مقابل النووية</code>\n\n"
        "✏️ <b>اكتب الآن 👇</b>",
        parse_mode='HTML'
    )


async def comp_no_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query   = update.callback_query
    await query.answer()
    user_id = query.from_user.id
    if user_id not in user_sessions:
        await query.edit_message_text("❌ الجلسة منتهية."); return
    if user_sessions[user_id].get("state") != "asking_comparison":
        await query.answer("هذا الزر لم يعد فعالاً.", show_alert=True); return
    session = user_sessions[user_id]
    session.pop("comparison_query", None)
    session["state"] = "choosing_file_format"
    await query.edit_message_text(
        "📁 <b>الآن اختر صيغة الملف الذي تريده:</b>",
        reply_markup=file_format_keyboard(), parse_mode='HTML'
    )


async def template_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query   = update.callback_query
    await query.answer()
    user_id = query.from_user.id
    tpl     = query.data.replace("tpl_", "")
    if user_id not in user_sessions:
        await query.edit_message_text("❌ الجلسة منتهية."); return
    if user_sessions[user_id].get("state") != "choosing_template":
        await query.answer("هذا الزر لم يعد فعالاً.", show_alert=True); return
    session = user_sessions[user_id]
    session["template"]    = tpl
    session["custom_mode"] = False
    session["state"]       = "choosing_file_format"
    await query.edit_message_text(
        "📁 <b>الآن اختر صيغة الملف الذي تريده:</b>",
        reply_markup=file_format_keyboard(), parse_mode='HTML'
    )


async def file_format_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query   = update.callback_query
    await query.answer()
    user_id = query.from_user.id
    fmt     = query.data.replace("format_", "")
    if user_id not in user_sessions:
        await query.edit_message_text("❌ الجلسة منتهية."); return
    session = user_sessions[user_id]
    if session.get("state") not in ["choosing_file_format", "in_queue"]:
        await query.answer("هذا الزر غير متاح الآن.", show_alert=True); return
    session["file_format"] = fmt
    session["state"] = "in_queue"
    pos = report_queue.qsize() + 1
    queue_positions[user_id] = pos
    await query.edit_message_text(build_queue_text(session, pos), parse_mode='HTML')
    await report_queue.put((user_id, session.copy(), query.message.message_id))


async def error_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    logger.error(f"Update error: {context.error}", exc_info=context.error)
    try:
        if update and update.effective_message:
            await update.effective_message.reply_text("❌ حدث خطأ. حاول مرة أخرى.")
    except Exception:
        pass


# ═══════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════
async def post_init(app):
    global report_queue
    report_queue = asyncio.Queue()
    asyncio.create_task(queue_worker(app))
    logger.info("✅ Queue worker started")


if __name__ == '__main__':
    flask_thread = threading.Thread(target=run_flask, daemon=True)
    flask_thread.start()
    logger.info("🌐 Flask started")

    token = os.getenv("TELEGRAM_TOKEN")
    if not token:
        logger.error("❌ TELEGRAM_TOKEN missing")
        exit(1)

    try:
        app = (
            ApplicationBuilder()
            .token(token)
            .post_init(post_init)
            .build()
        )
        app.add_handler(CommandHandler('start',  start))
        app.add_handler(CommandHandler('cancel', cancel))
        app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_message))

        app.add_handler(CallbackQueryHandler(title_auto_callback, pattern=r'^title_auto$'))
        app.add_handler(CallbackQueryHandler(language_callback,   pattern=r'^lang_'))
        app.add_handler(CallbackQueryHandler(depth_callback,      pattern=r'^depth_'))
        app.add_handler(CallbackQueryHandler(style_mode_callback, pattern=r'^style_'))
        app.add_handler(CallbackQueryHandler(template_callback,   pattern=r'^tpl_'))
        app.add_handler(CallbackQueryHandler(font_size_callback,  pattern=r'^fsize_'))
        app.add_handler(CallbackQueryHandler(font_callback,       pattern=r'^cfont_'))
        app.add_handler(CallbackQueryHandler(colors_callback,     pattern=r'^color_'))
        app.add_handler(CallbackQueryHandler(line_height_callback, pattern=r'^lh_'))
        app.add_handler(CallbackQueryHandler(page_margin_callback, pattern=r'^pm_'))
        app.add_handler(CallbackQueryHandler(header_style_callback, pattern=r'^hs_'))
        app.add_handler(CallbackQueryHandler(show_header_callback, pattern=r'^sh_'))
        app.add_handler(CallbackQueryHandler(comp_yes_callback,   pattern=r'^comp_yes$'))
        app.add_handler(CallbackQueryHandler(comp_no_callback,    pattern=r'^comp_no$'))
        app.add_handler(CallbackQueryHandler(file_format_callback, pattern=r'^format_(pdf|word)$'))
        app.add_error_handler(error_handler)

        logger.info("👻 Repooreto Bot v5.2 Ready!")
        print("=" * 60)
        print("👻 Repooreto — Smart University Reports Bot v5.2")
        print("=" * 60)
        app.run_polling(allowed_updates=Update.ALL_TYPES)

    except Exception as e:
        logger.error(f"❌ Startup failed: {e}", exc_info=True)
        exit(1)
